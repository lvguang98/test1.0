#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工伤案件管理系统 - 主程序
"""
import json
import os
import re
import shutil
import sys
import tempfile
import threading
import time
import traceback
from datetime import datetime
from PyQt5.QtCore import QSettings, Qt
from PyQt5.QtGui import QColor
from PyQt5.QtWidgets import (QApplication, QButtonGroup, QCompleter, QDialog,
                              QHBoxLayout, QLabel, QMainWindow, QMessageBox,
                              QPushButton, QRadioButton, QVBoxLayout)
from PyQt5.uic import loadUi
from docx import Document
from docxtpl import DocxTemplate
from openpyxl import Workbook, load_workbook
from config_manager import ConfigManager
from ctypes import *


class MainWindow(QMainWindow):
    # 人员类型常量
    PERSON_SELF = "本人"
    PERSON_WITNESS = "证人"
    PERSON_LEGAL = "法人"

    # 案件类型常量
    CASE_NORMAL = "普通案件"
    CASE_PERSONAL = "个人案件"
    CASE_DEATH = "死亡案件"
    CASE_PERSONAL_DEATH = "个人申请死亡案件"

    # 案本号前缀
    CASE_PREFIX_MAP = {
        CASE_NORMAL: "GS",
        CASE_PERSONAL: "GR",
        CASE_DEATH: "GSW",
        CASE_PERSONAL_DEATH: "GRW",
    }

    # ComboBox 默认占位文本
    PLACEHOLDER_EMPLOYER = "用人单位名称汇总"
    PLACEHOLDER_WORK_UNIT = "用工单位名称汇总"
    PLACEHOLDER_WORKPLACE = "工作场所名称汇总"

    def __init__(self):
        super().__init__()

        # 1. 加载界面
        loadUi("main_window.ui", self)
        self.setWindowTitle("工伤案件管理系统")

        # 2. 初始化配置管理器
        self.config = ConfigManager()

        self.BASE_DIR = os.path.dirname(__file__)
        self.TEMPLATE_DIR = os.path.join(self.BASE_DIR, "templates")

        # 3. 加载Excel数据到ComboBox
        self.load_excel_to_combobox()

        # 3.1 设置ComboBox的自动完成和失去焦点保存功能
        self.setup_combobox_autosave()

        # 3.2 连接删除按钮
        self.setup_delete_buttons()

        # 3.3 连接新增按钮放在这里
        self.setup_document_buttons()

        # 4. 加载保存的配置
        self.load_config()

        # 5. 连接信号
        self.checkBox_remember.stateChanged.connect(self.on_remember_changed)
        self.btn_generate_record.clicked.connect(self.on_generate_record)
        # 身份证号框失去焦点
        self.lineEdit_id_card.editingFinished.connect(self.auto_calculate_id_info)
        # 用人单位失去焦点时触发
        self.comboBox_employer.lineEdit().editingFinished.connect(self.auto_fill_applicant)
        # 本人姓名失去焦点时触发
        self.lineEdit_name.editingFinished.connect(self.auto_fill_applicant)
        # 添加这行 - 身份证读卡器按钮
        self.btn_read_id_card.clicked.connect(self.on_read_id_card)  # ← 新增

        # 6. 根据记住状态更新界面
        self.update_ui()

        # 连接人员类型切换信号
        self.radio_self.toggled.connect(self.on_person_type_changed)
        self.radio_witness.toggled.connect(self.on_person_type_changed)
        self.radio_legal_entity.toggled.connect(self.on_person_type_changed)

        # 连接案件类型复选框的信号
        self.checkBox_personal.stateChanged.connect(self.on_case_type_changed)
        self.checkBox_death.stateChanged.connect(self.on_case_type_changed)

        # 用人单位改变时更新申请人
        self.comboBox_employer.currentTextChanged.connect(self.on_case_type_changed)
        # 本人姓名改变时更新申请人
        self.lineEdit_name.textChanged.connect(self.on_case_type_changed)

        self.current_case_number = None  # 当前使用的案本号
        self.current_folder_path = None  # 当前使用的文件夹路径

        self.label_current_case.setText("当前案本：无")

    def on_read_id_card(self):
        """读取身份证按钮点击事件"""
        try:
            dll = windll.LoadLibrary("./sdtapi.dll")
            port = c_int32(1001)
            ifopen = c_int32(1)
            pucManaInfo = create_string_buffer(4)
            pucManaMsg = create_string_buffer(8)
            dll.SDT_StartFindIDCard(port, pucManaInfo, ifopen)
            dll.SDT_SelectIDCard(port, pucManaMsg, ifopen)
            pucCHMsg = create_unicode_buffer(256)
            pucPHMsg = create_string_buffer(1024)
            puiCHMsgLen = c_uint(0)
            puiPHMsgLen = c_uint(0)
            ret = dll.SDT_ReadBaseMsg(port, pucCHMsg, byref(puiCHMsgLen), pucPHMsg,
                                      byref(puiPHMsgLen), ifopen)
            if ret == 65:
                return
            dll.SDT_ClosePort(port)
            self.set_data('当前时期', datetime.now().strftime("%Y年%m月%d日%H时%M分"), 'output')

            role = self.get_current_role_type()
            self.process_id(pucCHMsg, role)

        except Exception as e:
            QMessageBox.warning(self, "错误", f"读取身份证失败: {str(e)}")

    def process_id(self, pucCHMsg, role):
        """处理身份证信息"""
        try:
            name = pucCHMsg.value[0:15].strip()
            self.data_model.update_basic_info(role, {'姓名': name})
            self.set_data(f"{role}姓名", name, 'basic')
            self.name_pane.setText(name)

            if len(pucCHMsg.value) >= 79:
                id_number = pucCHMsg.value[61:79].strip()
                self.data_model.update_basic_info(role, {'身份证号': id_number})
                self.set_data(f"{role}身份证号", id_number, 'basic')
                self.idnumer_pane.setText(id_number)

            if len(pucCHMsg.value) >= 61:
                address = pucCHMsg.value[26:61].strip()
                self.data_model.update_basic_info(role, {'身份证地址': address})
                self.set_data(f"{role}身份证地址", address, 'basic')
                self.textEdit.setText(address)

            self.process_id_info(role)
            self.calculate_age_from_id(role)

        except Exception as e:
            traceback.print_exc()

    def _load_excel_data(self, filepath):
        """从Excel文件加载数据到列表"""
        data_list = []
        try:
            if os.path.exists(filepath):
                wb = load_workbook(filepath)
                ws = wb.active
                for row in ws.iter_rows(min_row=1, max_col=1, values_only=True):
                    if row[0] and str(row[0]).strip():
                        data_list.append(str(row[0]).strip())
        except Exception as e:
            print(f"读取Excel失败 {filepath}: {e}")
        return data_list

    def _save_to_excel(self, filepath, new_item, column_name="汇总表"):
        """保存新项目到Excel文件"""
        try:
            if os.path.exists(filepath):
                wb = load_workbook(filepath)
                ws = wb.active
                # 找到第一个空行
                row = 1
                while ws.cell(row=row, column=1).value is not None:
                    row += 1
                ws.cell(row=row, column=1, value=new_item)
            else:
                # 文件不存在，创建新文件
                wb = Workbook()
                ws = wb.active
                ws.title = column_name
                ws.cell(row=1, column=1, value=column_name)
                ws.cell(row=2, column=1, value=new_item)

            wb.save(filepath)
            return True
        except Exception as e:
            print(f"保存到Excel失败: {e}")
            return False

    def _delete_from_excel(self, filepath, item_to_delete):
        """从Excel文件中删除指定项目"""
        try:

            if not os.path.exists(filepath):
                return False

            wb = load_workbook(filepath)
            ws = wb.active

            # 找到要删除的行（遍历直到连续遇到空行）
            row_to_delete = None
            empty_streak = 0
            for row in range(1, ws.max_row + 100):
                cell_value = ws.cell(row=row, column=1).value
                if cell_value is None:
                    empty_streak += 1
                    if empty_streak >= 5:  # 连续5个空行视为结束
                        break
                else:
                    empty_streak = 0
                    if str(cell_value).strip() == item_to_delete:
                        row_to_delete = row
                        break

            # 删除行
            if row_to_delete:
                ws.delete_rows(row_to_delete)
                wb.save(filepath)
                return True

            return False
        except Exception as e:
            print(f"从Excel删除失败: {e}")
            return False

    def _read_case_index(self):
        """读取案件索引文件"""
        index_file = os.path.join(self.BASE_DIR, "cases_index.json")
        try:
            if os.path.exists(index_file):
                with open(index_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                return {'cases': [], 'total_cases': 0, 'last_update': ''}
        except Exception as e:
            print(f"读取索引失败: {e}")
            return {'cases': [], 'total_cases': 0, 'last_update': ''}

    def _write_case_index(self, index_data):
        """写入案件索引文件原子操作"""
        index_file = os.path.join(self.BASE_DIR, "cases_index.json")
        try:

            # 使用临时文件
            with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', delete=False) as tf:
                json.dump(index_data, tf, ensure_ascii=False, indent=2)
                temp_file = tf.name

            # 替换原文件
            shutil.move(temp_file, index_file)
            return True

        except Exception as e:
            print(f"写入索引失败: {e}")
            # 注意：这里绝对不能调用任何可能再次调用本方法的方法
            return False

    def _find_case_in_index(self, case_number):
        """在索引中查找指定案本"""
        index_data = self._read_case_index()
        for case in index_data.get('cases', []):
            if case['case_number'] == case_number:
                return case, index_data
        return None, index_data

    def _get_approval_data(self, case_data):
        """获取审批表数据合并重复代码"""
        case_folder = os.path.join(self.BASE_DIR, case_data.get('folder_path', ''))

        # 查找审批表文件
        approval_file = None
        for file in os.listdir(case_folder):
            if file.endswith('_案件审批表.docx'):
                approval_file = os.path.join(case_folder, file)
                break

        if not approval_file:
            QMessageBox.warning(self, "提示", "请先生成案件审批表")
            return None

        # 从审批表读取数据
        approval_doc = Document(approval_file)

        申请时间 = ""
        受理时间 = ""
        综合情况 = ""
        医疗结论 = ""

        for table in approval_doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    text = cell.text
                    if "申请时间" in text and i + 1 < len(row.cells):
                        申请时间 = row.cells[i + 1].text.strip()
                    elif "受理时间" in text and i + 1 < len(row.cells):
                        受理时间 = row.cells[i + 1].text.strip()
                    elif "受伤经过" in text and i + 1 < len(row.cells):
                        综合情况 = row.cells[i + 1].text.strip()
                    elif "医疗诊断" in text and i + 1 < len(row.cells):
                        医疗结论 = row.cells[i + 1].text.strip()

        # 格式化时间调用已有方法
        申请时间 = self.format_date(申请时间)
        受理时间 = self.format_date(受理时间)

        return {
            '申请时间': 申请时间,
            '受理时间': 受理时间,
            '综合情况': 综合情况,
            '医疗结论': 医疗结论,
            'case_folder': case_folder,
            'approval_file': approval_file
        }

    def on_case_type_changed(self):
        """案件类型改变时，重新计算并更新申请人"""
        # 只处理本人类型
        if not self.radio_self.isChecked():
            return

        case_type = self.check_case_type()

        # 根据案件类型重新计算申请人
        if case_type in [self.CASE_NORMAL, self.CASE_DEATH]:
            # 普通案件或死亡案件：申请人 = 用人单位
            employer = self.comboBox_employer.currentText().strip()
            if employer and employer != self.PLACEHOLDER_EMPLOYER:
                self.lineEdit_applicant.setText(employer)
            else:
                self.lineEdit_applicant.clear()  # 如果没有用人单位，清空申请人

        elif case_type == self.CASE_PERSONAL:
            # 个人案件：申请人 = 本人姓名
            name = self.lineEdit_name.text().strip()
            if name:
                self.lineEdit_applicant.setText(name)
            else:
                self.lineEdit_applicant.clear()

        elif case_type == self.CASE_PERSONAL_DEATH:
            # 个人死亡案件：不清空，但可以给个提示
            if not self.lineEdit_applicant.text().strip():
                self.lineEdit_applicant.setPlaceholderText("请输入家属姓名")

    def auto_fill_applicant(self):
        """根据案件类型自动填充申请人"""
        # 只处理本人类型
        if not self.radio_self.isChecked():
            return

        case_type = self.check_case_type()

        # 普通案件或普通死亡案件：申请人 = 用人单位
        if case_type in [self.CASE_NORMAL, self.CASE_DEATH]:
            employer = self.comboBox_employer.currentText().strip()
            if employer and employer != self.PLACEHOLDER_EMPLOYER:
                self.lineEdit_applicant.setText(employer)

        # 个人案件：申请人 = 本人姓名
        elif case_type == self.CASE_PERSONAL:
            name = self.lineEdit_name.text().strip()
            if name:
                self.lineEdit_applicant.setText(name)

        # 个人死亡案件：不清空，让用户手动输入
        # elif case_type == self.CASE_PERSONAL_DEATH:
        #     pass  # 不做自动填充

    def setup_document_buttons(self):
        """连接各类文书生成按钮"""
        # 案件审批表
        self.btn_case_approval.clicked.connect(self.generate_case_approval)

        # 工伤告知书
        self.btn_injury_notice.clicked.connect(self.generate_injury_notice)

        # 谈话通知书
        self.btn_interview_notice.clicked.connect(self.generate_interview_notice)

    def _generate_document(self, template_name, filename, build_render_data,
                           require_approval=False, post_process=None):
        """通用文书生成：查案本→加载模板→渲染→保存→打开"""
        if not self.current_case_number:
            QMessageBox.warning(self, "错误", "请先生成本人案本或关联已有案本")
            return

        try:
            case_data, _ = self._find_case_in_index(self.current_case_number)
            if not case_data:
                QMessageBox.warning(self, "错误", f"未找到案本 {self.current_case_number} 的数据")
                return

            approval_data = None
            if require_approval:
                approval_data = self._get_approval_data(case_data)
                if not approval_data:
                    return
                if not approval_data['申请时间'] or not approval_data['受理时间']:
                    QMessageBox.warning(self, "提示", "审批表中未找到申请时间或受理时间")
                    return

            template_path = os.path.join(self.TEMPLATE_DIR, template_name)
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "错误", "模板不存在")
                return

            doc = DocxTemplate(template_path)
            render_data = build_render_data(case_data, approval_data)
            doc.render(render_data)

            case_folder = (approval_data['case_folder'] if approval_data
                           else os.path.join(self.BASE_DIR, case_data.get('folder_path', '')))
            filepath = os.path.join(case_folder, filename)
            doc.save(filepath)
            os.startfile(filepath)

            if post_process:
                post_process(filepath, case_data)

            self.statusBar().showMessage(f"已生成{filename}", 3000)

        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成失败: {str(e)}")
            traceback.print_exc()

    def generate_case_approval(self):
        """生成案件审批表"""
        injured_name = None

        def build_data(case_data, _):
            nonlocal injured_name
            injured_name = case_data.get('person_name', '')
            person_info = case_data.get('person_info', {})

            def process_self_intro(text):
                if not text:
                    return text
                return text[2:] if text.startswith("我是") else text

            def process_text(text):
                if not text:
                    return text
                text = text.replace("我们", "他们")
                text = text.replace("我", injured_name)
                return text

            def process_conclusion(text):
                if not text:
                    return text
                colon_index = text.find("：")
                if colon_index != -1:
                    return text[colon_index + 1:].strip()
                return text

            return {
                '案本号': case_data.get('case_number', ''),
                '受伤职工': injured_name,
                '申请人': case_data.get('applicant', ''),
                '性别': person_info.get('gender', ''),
                '年龄': person_info.get('age', ''),
                '身份证号': person_info.get('id_card', ''),
                '身份证地址': person_info.get('address', ''),
                '现住址': person_info.get('current_address', ''),
                '联系电话': person_info.get('phone', ''),
                '岗位': person_info.get('position', ''),
                '自我介绍': process_self_intro(person_info.get('自我介绍', '')),
                '受伤经过': process_text(person_info.get('受伤经过', '')),
                '就医情况': process_text(person_info.get('就医情况', '')),
                '医疗结论': process_conclusion(person_info.get('医疗结论', '')),
                '用人单位': case_data.get('employer', ''),
                '用工单位': case_data.get('work_unit', ''),
                '工作场所': case_data.get('workplace', ''),
                '条例': case_data.get('regulation', ''),
                '案件类型': case_data.get('case_type', ''),
                '操作员': case_data.get('operator', ''),
                '当前日期': datetime.now().strftime('%Y年%m月%d日'),
            }

        def post_process(filepath, case_data):
            approval_doc = Document(filepath)
            self.extract_approval_times(approval_doc, case_data)

        self._generate_document(
            template_name="工伤案件审批表模板.docx",
            filename=f"{self.current_case_number}_案件审批表.docx",
            build_render_data=build_data,
            post_process=post_process
        )

    def extract_approval_times(self, doc, case_data):
        """从审批表文档中提取申请时间和受理时间并保存到JSON"""
        try:
    
            application_time = ""
            acceptance_time = ""

            # 遍历表格查找申请时间和受理时间
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text.strip()

                    # 查找包含"申请时间"的行
                    if "申请时间" in row_text:
                        cells = row.cells
                        for i, cell in enumerate(cells):
                            if "申请时间" in cell.text:
                                if i + 1 < len(cells):
                                    time_value = cells[i + 1].text.strip()
                                    if time_value:
                                        application_time = self.format_date(time_value)
                                break

                    # 查找包含"受理时间"的行
                    if "受理时间" in row_text:
                        cells = row.cells
                        for i, cell in enumerate(cells):
                            if "受理时间" in cell.text:
                                if i + 1 < len(cells):
                                    time_value = cells[i + 1].text.strip()
                                    if time_value:
                                        acceptance_time = self.format_date(time_value)
                                break

            # 如果找到了时间，更新索引文件
            if application_time or acceptance_time:
                # 读取索引
                index_data = self._read_case_index()

                # 查找当前案件并更新时间
                for case in index_data.get('cases', []):
                    if case['case_number'] == case_data['case_number']:
                        if 'approval_info' not in case:
                            case['approval_info'] = {}

                        if application_time:
                            case['approval_info']['申请时间'] = application_time
                        if acceptance_time:
                            case['approval_info']['受理时间'] = acceptance_time

                        print(f"已保存申请时间: {application_time}, 受理时间: {acceptance_time}")
                        break

                # 保存更新后的索引
                self._write_case_index(index_data)

        except Exception as e:
            print(f"提取申请/受理时间失败: {e}")
            traceback.print_exc()

    def format_date(self, date_str):
        """将日期字符串格式化为 xxxx年xx月xx日"""
        if not date_str:
            return ""

        # 移除所有非数字字符
        digits = re.sub(r'\D', '', date_str)

        # 如果是8位数字如20260101
        if len(digits) == 8:
            year = digits[0:4]
            month = digits[4:6].lstrip('0')  # 去掉前导零
            day = digits[6:8].lstrip('0')  # 去掉前导零
            return f"{year}年{month}月{day}日"

        # 如果是其他格式，尝试解析
        try:
            # 尝试常见格式
            for fmt in ["%Y%m%d", "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"]:
                try:
                    dt = datetime.strptime(date_str, fmt)
                    return dt.strftime("%Y年%m月%d日")
                except:
                    continue
        except:
            pass

        # 如果都无法解析，返回原字符串
        return date_str

    def generate_injury_notice(self):
        """生成工伤认定告知书"""
        def build_data(case_data, approval_data):
            return {
                '用人单位': case_data.get('employer', ''),
                '申请人': case_data.get('applicant', ''),
                '申请时间': approval_data['申请时间'],
                '受理时间': approval_data['受理时间'],
                '受伤职工': case_data.get('person_name', ''),
                '综合情况': approval_data['综合情况'],
                '医疗结论': approval_data['医疗结论'],
                '条例': case_data.get('regulation', ''),
                '当前时期': datetime.now().strftime('%Y年%m月%d日'),
            }

        self._generate_document(
            template_name="工伤认定告知书模板.docx",
            filename=f"{self.current_case_number}_工伤认定告知书.docx",
            build_render_data=build_data,
            require_approval=True
        )

    def generate_interview_notice(self):
        """生成接受谈话通知书"""
        def build_data(case_data, approval_data):
            return {
                '用人单位': case_data.get('employer', ''),
                '申请人': case_data.get('applicant', ''),
                '申请时间': approval_data['申请时间'],
                '受理时间': approval_data['受理时间'],
                '本人姓名': case_data.get('person_name', ''),
                '本人身份证': case_data.get('person_info', {}).get('id_card', ''),
                '综合情况': approval_data['综合情况'],
                '医疗结论': approval_data['医疗结论'],
                '当前时期': datetime.now().strftime('%Y年%m月%d日'),
            }

        self._generate_document(
            template_name="接受谈话通知书模板.docx",
            filename=f"{self.current_case_number}_接受谈话通知书.docx",
            build_render_data=build_data,
            require_approval=True
        )

    def on_person_type_changed(self):
        """人员类型切换时的处理"""
        if self.sender().isChecked():
            person_type = self.check_person_type()
            self.statusBar().showMessage(f"当前人员类型: {person_type}", 1500)

            # 切换时清空相关字段
            if person_type in [self.PERSON_SELF, self.PERSON_WITNESS, self.PERSON_LEGAL]:
                self.clear_person_fields()

    def clear_person_fields(self):
        """清空人员信息字段"""
        self.lineEdit_name.clear()
        self.lineEdit_age.clear()
        self.comboBox_gender.setCurrentIndex(-1)
        self.lineEdit_id_card.clear()
        self.lineEdit_id_address.clear()
        self.lineEdit_current_address.clear()
        self.lineEdit_phone.clear()
        self.lineEdit_position.clear()

    def setup_delete_buttons(self):
        """设置删除按钮功能"""
        self.btn_delete_employer.clicked.connect(
            lambda: self.delete_from_excel('comboBox_employer', self.employer_list, "用人单位名称汇总.xlsx", "用人单位")
        )
        self.btn_delete_work_unit.clicked.connect(
            lambda: self.delete_from_excel('comboBox_work_unit', self.work_unit_list, "用工单位名称汇总.xlsx", "用工单位")
        )
        self.btn_delete_workplace.clicked.connect(
            lambda: self.delete_from_excel('comboBox_workplace', self.workplace_list, "工作场所名称汇总.xlsx", "工作场所")
        )

    def delete_from_excel(self, combobox_name, data_list, filename, column_name):
        """从Excel删除当前选中的项目"""
        # 获取对应的ComboBox
        combobox = getattr(self, combobox_name)

        # 获取当前选中的文本
        selected_text = combobox.currentText().strip()

        if not selected_text:
            self.statusBar().showMessage("请先选择要删除的项目", 2000)
            return

        # 确认对话框
        reply = QMessageBox.question(
            self, '确认删除',
            f'确定要删除 "{selected_text}" 吗？',
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.No:
            return

        try:
            # 1. 从内存列表中删除
            if selected_text in data_list:
                data_list.remove(selected_text)

            # 2. 从ComboBox中删除
            index = combobox.findText(selected_text)
            if index >= 0:
                combobox.removeItem(index)

            # 3. 从Excel文件中删除
            current_dir = os.path.dirname(os.path.abspath(__file__))
            filepath = os.path.join(current_dir, filename)

            if self._delete_from_excel(filepath, selected_text):
                self.statusBar().showMessage(f'已删除: {selected_text}', 3000)
            else:
                self.statusBar().showMessage("未在Excel中找到该项目", 3000)

            # 4. 清空当前选择
            combobox.setCurrentIndex(-1)
            combobox.setCurrentText("")

        except Exception as e:
            self.statusBar().showMessage(f"删除失败: {str(e)}", 3000)

    def setup_combobox_autosave(self):
        """设置ComboBox的自动完成和失去焦点保存功能"""
        # 为每个ComboBox设置相同的功能
        for combobox_name in ['comboBox_employer', 'comboBox_work_unit', 'comboBox_workplace']:
            combobox = getattr(self, combobox_name)

            # 设置可编辑
            combobox.setEditable(True)

            # 设置自动完成，显示最多3个相似项

            # 获取当前列表数据
            if combobox_name == 'comboBox_employer':
                data_list = self.employer_list
            elif combobox_name == 'comboBox_work_unit':
                data_list = self.work_unit_list
            else:  # comboBox_workplace
                data_list = self.workplace_list

            # 创建自动完成器
            completer = QCompleter(data_list)
            completer.setFilterMode(Qt.MatchContains)  # 包含匹配
            completer.setMaxVisibleItems(3)  # 最多显示3个
            combobox.setCompleter(completer)

            # 获取ComboBox内部的QLineEdit并连接失去焦点事件
            line_edit = combobox.lineEdit()
            line_edit.editingFinished.connect(
                lambda le=line_edit, cb=combobox, name=combobox_name, lst=data_list:
                self.on_combobox_editing_finished(le, cb, name, lst)
            )

    def on_combobox_editing_finished(self, line_edit, combobox, combobox_name, current_list):
        """ComboBox失去焦点时的处理"""
        # 获取用户输入的文本
        user_input = line_edit.text().strip()

        if not user_input:
            return  # 如果输入为空，不处理

        # 检查是否已经在列表中
        if user_input in current_list:
            return  # 如果已经在列表中，不重复添加

        # 如果不在列表中，保存到Excel
        self.save_to_excel(combobox_name, user_input, current_list)

        # 添加到内存列表和ComboBox
        current_list.append(user_input)
        combobox.addItem(user_input)

        # 保持用户输入的内容显示在界面上
        combobox.setCurrentText(user_input)

    def save_to_excel(self, combobox_name, new_item, current_list):
        """保存新项目到对应的Excel文件"""
        # 确定文件名和列名
        if combobox_name == 'comboBox_employer':
            filename = "用人单位名称汇总.xlsx"
            column_name = "用人单位"
        elif combobox_name == 'comboBox_work_unit':
            filename = "用工单位名称汇总.xlsx"
            column_name = "用工单位"
        else:  # comboBox_workplace
            filename = "工作场所名称汇总.xlsx"
            column_name = "工作场所"

        current_dir = os.path.dirname(os.path.abspath(__file__))
        filepath = os.path.join(current_dir, filename)

        # 调用统一方法保存
        return self._save_to_excel(filepath, new_item, column_name)

    def auto_calculate_id_info(self):
        """自动计算身份证信息"""
        id_card = self.lineEdit_id_card.text().strip()
        if id_card:
            _, age, gender = self.calculate_id_info(id_card)
            if age:
                self.lineEdit_age.setText(str(age))
            if gender:
                self.comboBox_gender.setCurrentText(gender)

    def load_excel_to_combobox(self):
        """从Excel文件加载数据到ComboBox"""
        current_dir = os.path.dirname(os.path.abspath(__file__))

        # 加载用人单位
        self.employer_list = self._load_excel_data(os.path.join(current_dir, "用人单位名称汇总.xlsx"))
        self.comboBox_employer.addItems(self.employer_list)

        # 加载用工单位
        self.work_unit_list = self._load_excel_data(os.path.join(current_dir, "用工单位名称汇总.xlsx"))
        self.comboBox_work_unit.addItems(self.work_unit_list)

        # 加载工作场所
        self.workplace_list = self._load_excel_data(os.path.join(current_dir, "工作场所名称汇总.xlsx"))
        self.comboBox_workplace.addItems(self.workplace_list)

    def load_config(self):
        """加载配置到界面"""
        config = self.config.load_config()

        # 设置控件内容
        self.lineEdit_operator.setText(config['operator'])
        self.lineEdit_api_url.setText(config['api_url'])
        self.lineEdit_api_key.setText(config['api_key'])
        self.checkBox_remember.setChecked(config['remember'])

    def update_ui(self):
        """更新界面状态"""
        remember = self.checkBox_remember.isChecked()

        # 设置输入框是否可编辑
        self.lineEdit_operator.setEnabled(not remember)
        self.lineEdit_api_url.setEnabled(not remember)
        self.lineEdit_api_key.setEnabled(not remember)

        # 设置样式
        if remember:
            style = "background-color: #f0f0f0; color: #666;"
            self.statusBar().showMessage("配置已记住，取消勾选可修改", 2000)
        else:
            style = ""

        self.lineEdit_operator.setStyleSheet(f"QLineEdit {{ {style} }}")
        self.lineEdit_api_url.setStyleSheet(f"QLineEdit {{ {style} }}")
        self.lineEdit_api_key.setStyleSheet(f"QLineEdit {{ {style} }}")

    def on_remember_changed(self):
        """记住我复选框状态变化"""
        remember = self.checkBox_remember.isChecked()

        if remember:
            # 保存当前配置
            operator = self.lineEdit_operator.text().strip()
            api_url = self.lineEdit_api_url.text().strip()
            api_key = self.lineEdit_api_key.text().strip()

            self.config.save_config(operator, api_url, api_key, True)
            self.statusBar().showMessage("配置已保存", 1500)
        else:
            # 清除配置
            self.config.clear_config()
            self.statusBar().showMessage("配置已清除", 1500)

        # 更新界面
        self.update_ui()

    def check_case_type(self):
        """检查案件类型"""
        is_personal = self.checkBox_personal.isChecked()
        is_death = self.checkBox_death.isChecked()

        if is_personal and is_death:
            return self.CASE_PERSONAL_DEATH
        elif is_personal:
            return self.CASE_PERSONAL
        elif is_death:
            return self.CASE_DEATH
        else:
            return self.CASE_NORMAL

    def check_person_type(self):
        """检查人员类型"""
        if self.radio_self.isChecked():
            return self.PERSON_SELF
        elif self.radio_witness.isChecked():
            return self.PERSON_WITNESS
        elif self.radio_legal_entity.isChecked():
            return self.PERSON_LEGAL

    def calculate_id_info(self, id_card):
        """根据身份证号计算年龄和性别"""
        if len(id_card) != 18:
            return id_card, None, None

        # 提取出生年月日
        birth_year = int(id_card[6:10])
        birth_month = int(id_card[10:12])
        birth_day = int(id_card[12:14])

        # 计算年龄
        current_year = datetime.now().year
        current_month = datetime.now().month
        current_day = datetime.now().day

        age = current_year - birth_year
        if current_month < birth_month or (current_month == birth_month and current_day < birth_day):
            age -= 1

        # 计算性别
        gender_num = int(id_card[16])
        gender = "男" if gender_num % 2 == 1 else "女"

        return id_card, age, gender

    def on_generate_record(self):
        """生成笔录按钮点击"""
        # 检查个人死亡案件的申请人
        if self.check_case_type() == self.CASE_PERSONAL_DEATH and self.radio_self.isChecked():
            if not self.lineEdit_applicant.text().strip():
                QMessageBox.warning(self, "提示", "请填写申请人信息家属姓名")
                return

        # 1. 收集数据
        data = self.collect_form_data()
        if not data:  # 如果收集数据失败比如申请人未填写
            return

        # 2. 根据人员类型分流
        if data['人员类型'] == self.PERSON_SELF:
            self.handle_person_case(data)
        elif data['人员类型'] == self.PERSON_WITNESS:
            self.handle_witness_case(data)
        elif data['人员类型'] == self.PERSON_LEGAL:
            self.handle_legal_case(data)

    def handle_person_case(self, data):
        # 1. 生成自我介绍
        description = self.generate_description(data)

        case_type = data['案件类型']

        # 2. 检查是否已有案本
        index_data = self._read_case_index()

        if index_data.get('cases'):

            # 搜索同名案件
            same_person_cases = []
            for case in index_data.get('cases', []):
                if case['person_name'] == data['受伤职工']:
                    same_person_cases.append(case)

            if same_person_cases:
                selected_case = self.show_case_selection_dialog(
                    data['受伤职工'],
                    same_person_cases,
                    data['本人身份证号']
                )

                if selected_case == "new":
                    pass
                elif selected_case:
                    self.current_case_number = selected_case['case_number']
                    self.current_folder_path = selected_case['folder_path']

                    # 更新案本号显示只改文本
                    self.label_current_case.setText(f"当前案本：{self.current_case_number}")

                    person_info = selected_case.get('person_info', {})
                    self.lineEdit_name.setText(person_info.get('name', ''))
                    self.lineEdit_id_card.setText(selected_case.get('id_card', ''))
                    self.lineEdit_phone.setText(person_info.get('phone', ''))
                    if selected_case.get('id_card'):
                        self.auto_calculate_id_info()

                    self.statusBar().showMessage(f"已关联案本: {selected_case['case_number']}", 3000)

                    if "死亡" in case_type:
                        QMessageBox.information(self, "提示",
                                                f"已关联死亡职工案本：{selected_case['case_number']}\n\n请继续输入证人笔录信息",
                                                QMessageBox.Ok)
                    return
                else:
                    return

        # ========== 新建案件 ==========
        case_number = self.generate_case_number(data['受伤职工'])
        data['案本号'] = case_number
        year_folder = self.get_current_year_folder()
        case_folder = os.path.join(year_folder, case_number)
        os.makedirs(case_folder, exist_ok=True)

        # 保存当前使用的案本信息
        self.current_case_number = case_number
        self.current_folder_path = f"{datetime.now().year}/{case_number}"

        # 更新案本号显示只改文本
        self.label_current_case.setText(f"当前案本：{self.current_case_number}")

        # 在数据中添加自我介绍
        data['自我介绍'] = description

        # 预留三个字段，等待后续提取
        data['受伤经过'] = ''
        data['就医情况'] = ''
        data['医疗结论'] = ''

        # 更新索引
        self.update_case_index(case_number, data['受伤职工'], data)

        # 判断是否为死亡案件
        if "死亡" in case_type:
            QMessageBox.information(self, "提示",
                                    f"死亡职工信息已保存\n案本号：{case_number}\n\n请继续输入证人笔录信息",
                                    QMessageBox.Ok)
        else:
            template_name = self.get_template_name(data)
            self.generate_transcript(case_folder, template_name, data)

    def handle_witness_case(self, data):
        self._handle_nonparty_case(data, self.PERSON_WITNESS)

    def handle_legal_case(self, data):
        try:
            self._handle_nonparty_case(data, self.PERSON_LEGAL)
        except Exception as e:
            traceback.print_exc()
            self.statusBar().showMessage(f"错误: {str(e)}", 3000)

    def _handle_nonparty_case(self, data, person_type):
        """处理证人/法人案件的通用方法"""
        description = self.generate_description(data)
        data['自我介绍'] = description

        if not self.current_case_number:
            QMessageBox.warning(self, "错误", "请先生成本人案本或关联已有案本")
            return

        person_name = data.get(f'{person_type}姓名', '')
        data['案本号'] = self.current_case_number

        year_folder = self.get_current_year_folder()
        case_folder = os.path.join(year_folder, self.current_case_number)
        if not os.path.exists(case_folder):
            os.makedirs(case_folder, exist_ok=True)

        existing_files = []
        if os.path.exists(case_folder):
            for file in os.listdir(case_folder):
                if file.endswith('.docx') and person_type in file:
                    existing_files.append(file)

        if not existing_files:
            template_name = self.get_template_name(data)
            self._create_nonparty_transcript(case_folder, data, person_type, number=1, template_name=template_name)
            return

        person_exists = False
        max_number = 0
        existing_file = None
        pattern = rf'{person_type}(\d+)_(.+?)\.docx'

        for file in existing_files:
            match = re.search(pattern, file)
            if match:
                num = int(match.group(1))
                existing_name = match.group(2)
                max_number = max(max_number, num)
                if existing_name == person_name:
                    person_exists = True
                    existing_file = os.path.join(case_folder, file)

        if person_exists:
            reply = QMessageBox.question(
                self, f'{person_type}已存在',
                f'{person_type} "{person_name}" 已有笔录\n是否打开？\n\n选"是"=打开\n选"否"=新建另一份',
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            if reply == QMessageBox.Yes:
                if existing_file and os.path.exists(existing_file):
                    os.startfile(existing_file)
                    self.statusBar().showMessage(f"已打开{person_type}笔录", 3000)
                else:
                    QMessageBox.warning(self, "错误", f"找不到{person_type}笔录文件")
            else:
                template_name = self.get_template_name(data)
                self._create_nonparty_transcript(case_folder, data, person_type, number=max_number + 1,
                                                 template_name=template_name)
        else:
            template_name = self.get_template_name(data)
            self._create_nonparty_transcript(case_folder, data, person_type, number=max_number + 1,
                                             template_name=template_name)

    def _create_nonparty_transcript(self, case_folder, data, person_type, number, template_name):
        """生成证人/法人笔录"""
        return self.generate_transcript_unified(
            case_folder=case_folder,
            data=data,
            template_name=template_name,
            file_prefix=person_type,
            person_type=person_type,
            person_name=data.get(f'{person_type}姓名', '')
        )

    def generate_transcript_unified(self, case_folder, data, template_name, file_prefix, person_type, person_name):
        """
        统一的笔录生成方法使用python-docx-template库
        """
        try:
            template_path = os.path.join(self.TEMPLATE_DIR, template_name)
            if not os.path.exists(template_path):
                self.statusBar().showMessage(f"模板不存在: {template_name}", 3000)
                return None

            # 1. 使用DocxTemplate加载模板
            doc = DocxTemplate(template_path)

            # 2. 生成自我介绍文本
            self_intro_text = self.generate_description(data)

            # 3. 准备所有替换数据
            render_data = {
                # 基本信息
                '受伤职工': data.get('受伤职工', ''),
                '用人单位': data.get('用人单位', ''),
                '用工单位': data.get('用工单位', ''),
                '工作场所': data.get('工作场所', ''),
                '操作员': data.get('操作员', ''),
                '当前日期': datetime.now().strftime('%Y年%m月%d日'),
                '当前时间': datetime.now().strftime('%H时%M分'),

                # 自我介绍 - 直接放入生成的文本
                '自我介绍': self_intro_text,

                # 人员特定信息
                f'{person_type}姓名': person_name,
                f'{person_type}性别': data.get(f'{person_type}性别', ''),
                f'{person_type}年龄': data.get(f'{person_type}年龄', ''),
                f'{person_type}身份证': data.get(f'{person_type}身份证号', ''),
                f'{person_type}身份证地址': data.get(f'{person_type}身份证地址', ''),
                f'{person_type}电话': data.get(f'{person_type}电话', ''),
                f'{person_type}岗位': data.get(f'{person_type}岗位', ''),
            }

            # 4. 渲染模板替换所有普通占位符
            doc.render(render_data)

            # 5. 保存临时文件
            temp_path = os.path.join(case_folder, "temp_render.docx")
            doc.save(temp_path)

            # 6. 用python-docx打开处理问答句如果需要
            final_doc = Document(temp_path)

            # 7. 添加问答句仅针对特殊案件
            case_type = data.get('案件类型', '')
            if case_type in [self.CASE_PERSONAL, self.CASE_DEATH, self.CASE_PERSONAL_DEATH]:
                self.add_questions_to_doc_with_format(final_doc, case_type, data)

            # 8. 生成最终文件名
            injured_name = data.get('受伤职工', '')
            max_num = 0
            if os.path.exists(case_folder):
                for file in os.listdir(case_folder):
                    pattern = rf'{injured_name}_{file_prefix}(\d+)_'
                    match = re.search(pattern, file)
                    if match:
                        num = int(match.group(1))
                        max_num = max(max_num, num)

            next_num = max_num + 1
            filename = f"{injured_name}_{file_prefix}{next_num:02d}_{person_name}.docx"
            filepath = os.path.join(case_folder, filename)

            # 9. 保存最终文件
            final_doc.save(filepath)

            # 10. 删除临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)

            # 11. 打开文件
            os.startfile(filepath)

            self.statusBar().showMessage(f"{person_type}笔录已生成: {filename}", 3000)

            # 12. 更新索引
            self.update_case_index(data.get('案本号', ''), data.get('受伤职工', ''), data)

            # 13. 如果是本人笔录，启动后台线程监控文件关闭
            if person_type == self.PERSON_SELF:
                def wait_for_file_close():
                    time.sleep(2)
                    MAX_WAIT = 300  # 最多等5分钟
                    elapsed = 0
                    while elapsed < MAX_WAIT:
                        try:
                            with open(filepath, 'a'):
                                pass
                            self.extract_person_info_from_doc(filepath, data.get('案本号', ''))
                            return
                        except (IOError, OSError):
                            time.sleep(2)
                            elapsed += 2
                    # 超时，放弃提取
                    self.statusBar().showMessage("笔录信息自动提取超时，可手动复制", 5000)

                threading.Thread(target=wait_for_file_close, daemon=True).start()

            return filepath

        except Exception as e:
            self.statusBar().showMessage(f"生成失败: {str(e)}", 3000)
            traceback.print_exc()
            return None

    def add_questions_to_doc_with_format(self, doc, case_type, data):
        """添加带格式的问答句仅用于特殊案件"""
        questions = self.generate_case_questions(case_type, data)
        if not questions:
            return

        # 添加一个空行分隔
        doc.add_paragraph()

        for q in questions:
            # 添加问句
            para = doc.add_paragraph()
            run = para.add_run(q)

            # 设置问句格式可以加粗
            if q.startswith("问："):
                run.font.bold = True

            # 添加一个空行分隔每个问答对
            doc.add_paragraph()

    def generate_case_questions(self, case_type, data):
        """根据案件类型生成对应的问答句"""
        if case_type == self.CASE_PERSONAL:
            return [
                "问：你是个人申请工伤认定吗？",
                "答：是的，我是个人申请。",
                "问：单位为什么没有为你申请？",
                "答：单位说让我自己申请。",
            ]
        elif case_type == self.CASE_DEATH:
            return [
                "问：你是死亡职工的家属吗？",
                "答：是的，我是他的家属。",
                "问：死亡时间和原因是什么？",
                "答：...",
            ]
        elif case_type == self.CASE_PERSONAL_DEATH:
            return [
                "问：你是以家属身份个人申请工亡吗？",
                "答：是的。",
                "问：单位没有为死者申报吗？",
                "答：没有。",
            ]
        else:
            return []

    def search_same_name_cases(self, name, id_card):
        """搜索同名案件"""
        cases = []

        # 读取索引文件
        index_data = self._read_case_index()

        for case in index_data.get('cases', []):
            if case['person_name'] == name:
                # 检查身份证号如果有
                case_id = case.get('id_card', '')
                if id_card and case_id:
                    if id_card == case_id:
                        case['match_type'] = '身份证完全匹配'
                    else:
                        case['match_type'] = '姓名匹配(身份证不同)'
                else:
                    case['match_type'] = '姓名匹配'

                cases.append(case)

        return cases

    def generate_description(self, data):
        """根据人员类型和单位情况生成描述语句"""
        person_type = data['人员类型']

        # 获取姓名
        if person_type == self.PERSON_SELF:
            name = data.get('本人姓名', '')
        elif person_type == self.PERSON_WITNESS:
            name = data.get('证人姓名', '')
        else:  # 法人
            name = data.get('法人姓名', '')

        employer = data.get('用人单位', '')
        work_unit = data.get('用工单位', '')
        workplace = data.get('工作场所', '')

        # 获取受伤职工姓名
        injured_name = data.get('受伤职工', '')

        # 从索引文件获取本人岗位
        injured_position = ""
        if self.current_case_number:
            index_file = os.path.join(self.BASE_DIR, "cases_index.json")
            if os.path.exists(index_file):
                try:
                    with open(index_file, 'r', encoding='utf-8') as f:
                        index_data = json.load(f)
                    for case in index_data.get('cases', []):
                        if case['case_number'] == self.current_case_number:
                            injured_position = case.get('person_info', {}).get('position', '')
                            break
                except:
                    pass

        has_employer = bool(employer)
        has_work_unit = bool(work_unit)
        has_workplace = bool(workplace)

        # 法人特殊处理
        if person_type == self.PERSON_LEGAL:
            # 第一部分：法人自身介绍
            if has_employer:
                description = f"我是{name}，是{employer}的法定代表人，负责公司全面管理工作。"
            else:
                description = f"我是{name}，是法定代表人，负责公司全面管理工作。"

            # 第二部分：介绍受伤职工
            if injured_name:
                if has_work_unit and has_workplace:
                    description += f"{injured_name}是我公司指派到{work_unit}承建的{workplace}工作的员工，从事{injured_position}工作。"
                elif has_work_unit:
                    description += f"{injured_name}是我公司指派到{work_unit}工作的员工，从事{injured_position}工作。"
                elif has_workplace:
                    description += f"{injured_name}是我公司指派到{workplace}工作的员工，从事{injured_position}工作。"
                else:
                    description += f"{injured_name}是我公司的员工，从事{injured_position}工作。"

            return description

        # 本人和证人的原有逻辑
        if person_type in [self.PERSON_SELF, self.PERSON_WITNESS]:
            position = data.get(f'{person_type}岗位', '')
            if has_employer and has_work_unit and has_workplace:
                description = f"我是{name}，系{employer}的职工，被指派到{work_unit}的{workplace}工作。从事{position}工作。"
            elif has_employer and has_work_unit:
                description = f"我是{name}，系{employer}的职工，被指派到{work_unit}工作。从事{position}工作。"
            elif has_employer and has_workplace:
                description = f"我是{name}，系{employer}的职工，被指派到{workplace}工作。从事{position}工作。"
            elif has_employer:
                description = f"我是{name}，系{employer}的职工。从事{position}工作。"
            else:
                description = f"我是{name}。从事{position}工作。"

        return description

    def show_case_selection_dialog(self, name, cases, id_card):
        """显示案件选择对话框支持红色显示身份证不同的案件"""

        dialog = QDialog(self)
        dialog.setWindowTitle("发现同名案件")
        dialog.resize(450, 350)

        layout = QVBoxLayout()

        # 标题
        if id_card:
            title = f'发现与"{name}"(身份证:{id_card[-4:]})同名的案件:'
        else:
            title = f'发现与"{name}"同名的案件:'
        layout.addWidget(QLabel(title))

        # 创建单选按钮组
        button_group = QButtonGroup()

        # 添加"新建案件"选项
        new_case_radio = QRadioButton("新建案件不关联已有")
        new_case_radio.setChecked(True)
        button_group.addButton(new_case_radio, 0)
        layout.addWidget(new_case_radio)

        layout.addWidget(QLabel("已有案本:"))

        # 添加已有案件选项
        for i, case in enumerate(cases, 1):
            case_num = case['case_number']
            # 从 person_info 里取身份证
            case_id = case.get('person_info', {}).get('id_card', '')

            # 创建水平布局放单选按钮和文本
            h_layout = QHBoxLayout()
            radio = QRadioButton()
            button_group.addButton(radio, i)
            h_layout.addWidget(radio)

            # 创建显示文本的标签
            if case_id:
                id_display = case_id[-4:] if len(case_id) >= 4 else case_id
                text = f"{case_num} (身份证:{id_display})"
            else:
                text = f"{case_num} (身份证:无)"

            label = QLabel(text)

            # 如果输入的身份证和案件的身份证不同，设为红色
            if id_card and case_id and id_card != case_id:
                label.setStyleSheet("color: red;")

            h_layout.addWidget(label)
            h_layout.addStretch()
            layout.addLayout(h_layout)

        # 按钮区域
        btn_layout = QHBoxLayout()
        btn_ok = QPushButton("确定")
        btn_cancel = QPushButton("取消")

        btn_ok.clicked.connect(dialog.accept)
        btn_cancel.clicked.connect(dialog.reject)

        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

        dialog.setLayout(layout)

        if dialog.exec_() == QDialog.Accepted:
            selected_id = button_group.checkedId()
            if selected_id == 0:
                return "new"
            elif selected_id > 0:
                return cases[selected_id - 1]

        return None

    def show_transcript_exists_dialog(self, case_number):
        """显示已有笔录对话框"""

        dialog = QDialog(self)
        dialog.setWindowTitle("已有本人笔录")
        dialog.resize(300, 150)

        layout = QVBoxLayout()
        layout.addWidget(QLabel(f"案本 {case_number} 已有本人笔录文件"))

        btn_layout = QHBoxLayout()

        btn_open = QPushButton("打开现有笔录")
        btn_supplement = QPushButton("生成补充笔录")
        btn_cancel = QPushButton("取消")

        btn_open.clicked.connect(lambda: dialog.done(1))
        btn_supplement.clicked.connect(lambda: dialog.done(2))
        btn_cancel.clicked.connect(dialog.reject)

        btn_layout.addWidget(btn_open)
        btn_layout.addWidget(btn_supplement)
        btn_layout.addWidget(btn_cancel)

        layout.addLayout(btn_layout)
        dialog.setLayout(layout)

        result = dialog.exec_()

        if result == 1:
            return "open"
        elif result == 2:
            return "supplement"
        else:
            return "cancel"

    def collect_form_data(self):
        """收集当前表单所有数据"""
        id_card = self.lineEdit_id_card.text().strip()
        _, age, gender = self.calculate_id_info(id_card) if id_card else (None, None, None)

        # 获取人员类型作为前缀
        prefix = self.check_person_type()

        # 获取条例选择
        regulation_text = self.comboBox_regulations.currentText().strip()

        # 简单的判断：如果ComboBox显示的是文件名，就视为空
        employer = self.comboBox_employer.currentText().strip()
        if employer == self.PLACEHOLDER_EMPLOYER:
            employer = ""

        work_unit = self.comboBox_work_unit.currentText().strip()
        if work_unit == self.PLACEHOLDER_WORK_UNIT:
            work_unit = ""

        workplace = self.comboBox_workplace.currentText().strip()
        if workplace == self.PLACEHOLDER_WORKPLACE:
            workplace = ""

        applicant = self.lineEdit_applicant.text().strip()

        # 如果有当前案本号，从索引文件获取本人姓名
        injured_worker = ""
        if self.current_case_number:
            # 从索引文件读取本人姓名
            index_file = os.path.join(self.BASE_DIR, "cases_index.json")
            if os.path.exists(index_file):
                try:
                    with open(index_file, 'r', encoding='utf-8') as f:
                        index_data = json.load(f)
                    for case in index_data.get('cases', []):
                        if case['case_number'] == self.current_case_number:
                            injured_worker = case.get('person_name', '')
                            break
                except:
                    pass

        # 如果没有当前案本新建本人案件时，才从输入框获取
        if not injured_worker and prefix == self.PERSON_SELF:
            injured_worker = self.lineEdit_name.text().strip()

        # 个人死亡案件检查
        if self.check_case_type() == self.CASE_PERSONAL_DEATH and self.radio_self.isChecked():
            if not applicant:
                return None

        data = {
            '案本号': '',
            '受伤职工': injured_worker,  # 现在始终是本人姓名
            '申请人': applicant,
            '用人单位': employer,
            '用工单位': work_unit,
            '工作场所': workplace,
            '人员类型': prefix,
            '案件类型': self.check_case_type(),
            '条例': regulation_text,
            '操作员': self.lineEdit_operator.text().strip(),
            '当前日期': datetime.now().strftime('%Y年%m月%d日'),
            '当前时间': datetime.now().strftime('%H时%M分'),
        }

        # 用变量作为前缀
        data[f'{prefix}姓名'] = self.lineEdit_name.text().strip()
        data[f'{prefix}性别'] = gender if gender else self.comboBox_gender.currentText()
        data[f'{prefix}年龄'] = str(age) if age else self.lineEdit_age.text().strip()
        data[f'{prefix}身份证号'] = id_card
        data[f'{prefix}身份证地址'] = self.lineEdit_id_address.text().strip()
        data[f'{prefix}现住址'] = self.lineEdit_current_address.text().strip()
        data[f'{prefix}电话'] = self.lineEdit_phone.text().strip()
        data[f'{prefix}岗位'] = self.lineEdit_position.text().strip()

        return data

    def generate_transcript(self, case_folder, template_name, data):
        """生成本人笔录"""
        return self.generate_transcript_unified(
            case_folder=case_folder,
            data=data,
            template_name=template_name,
            file_prefix=self.PERSON_SELF,
            person_type=self.PERSON_SELF,
            person_name=data.get('本人姓名', '')
        )

    def extract_person_info_from_doc(self, doc_file, case_number):
        """从Word文档中提取本人关键信息"""
        try:

            if not os.path.exists(doc_file):
                return

            doc = Document(doc_file)

            # 要搜索的关键词
            question_keywords = {
                '受伤经过': ['什么工作原因', '事故发生', '具体经过', '详细描述', '日常接触'],
                '就医情况': ['受伤后', '哪个医院', '是谁送你', '何处就诊', '哪些症状'],
                '医疗结论': ['此次受伤', '医院对你', '医疗结论', '诊断结论', '医院最终']
            }

            extracted_info = {}

            # 遍历所有段落，查找问题
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue

                for info_key, keywords in question_keywords.items():
                    if info_key in extracted_info:
                        continue

                    match_count = 0
                    for keyword in keywords:
                        if keyword in text:
                            match_count += 1

                    if match_count >= 2:
                        if i + 1 < len(doc.paragraphs):
                            answer = doc.paragraphs[i + 1].text.strip()
                            if answer.startswith('答：'):
                                answer = answer[2:].strip()
                            elif answer.startswith('答:'):
                                answer = answer[1:].strip()

                            extracted_info[info_key] = answer
                            break

            if extracted_info:
                self.update_extracted_info_in_index(case_number, extracted_info)

        except Exception as e:
            print(f"提取信息失败: {e}")

    def update_extracted_info_in_index(self, case_number, extracted_info):
        """只更新受伤经过、就医情况、医疗结论三个字段"""
        try:
            # 读取现有索引
            index_data = self._read_case_index()

            for case in index_data.get('cases', []):
                if case['case_number'] == case_number:
                    if 'person_info' in case:
                        case['person_info']['受伤经过'] = extracted_info.get('受伤经过', '')
                        case['person_info']['就医情况'] = extracted_info.get('就医情况', '')
                        case['person_info']['医疗结论'] = extracted_info.get('医疗结论', '')
                    break

            # 写入文件
            self._write_case_index(index_data)

        except Exception as e:
            print(f"更新提取信息失败: {e}")

    def update_case_index(self, case_number, person_name, data):
        """更新案件索引文件合并数据，避免覆盖"""
        try:

            # 读取现有索引
            index_data = self._read_case_index()

            # 查找现有案件
            found = False
            for i, existing_case in enumerate(index_data['cases']):
                if existing_case['case_number'] == case_number:
                    # 获取现有的 person_info
                    existing_person_info = existing_case.get('person_info', {})

                    # 构建新的 person_info保留旧数据，用新数据覆盖
                    new_person_info = {
                        'name': data.get('本人姓名', existing_person_info.get('name', '')),
                        'gender': data.get('本人性别', existing_person_info.get('gender', '')),
                        'age': data.get('本人年龄', existing_person_info.get('age', '')),
                        'phone': data.get('本人电话', existing_person_info.get('phone', '')),
                        'id_card': data.get('本人身份证号', existing_person_info.get('id_card', '')),
                        'address': data.get('本人身份证地址', existing_person_info.get('address', '')),
                        'current_address': data.get('本人现住址', existing_person_info.get('current_address', '')),
                        'position': data.get('本人岗位', existing_person_info.get('position', '')),
                        '自我介绍': data.get('自我介绍', existing_person_info.get('自我介绍', '')),
                        '受伤经过': data.get('受伤经过', existing_person_info.get('受伤经过', '')),
                        '就医情况': data.get('就医情况', existing_person_info.get('就医情况', '')),
                        '医疗结论': data.get('医疗结论', existing_person_info.get('医疗结论', ''))
                    }

                    # 构建完整的案件数据
                    index_data['cases'][i] = {
                        'case_number': case_number,
                        'person_name': person_name,
                        'applicant': data.get('申请人', ''),
                        'case_type': data.get('案件类型', existing_case.get('case_type', '')),
                        'year': datetime.now().year,
                        'folder_path': existing_case.get('folder_path', f"{datetime.now().year}/{case_number}"),
                        'created_date': existing_case.get('created_date', datetime.now().strftime('%Y-%m-%d')),
                        'employer': data.get('用人单位', existing_case.get('employer', '')),
                        'work_unit': data.get('用工单位', existing_case.get('work_unit', '')),
                        'workplace': data.get('工作场所', existing_case.get('workplace', '')),
                        'regulation': data.get('条例', existing_case.get('regulation', '')),
                        'operator': data.get('操作员', existing_case.get('operator', '')),
                        'person_info': new_person_info,
                        'witnesses': existing_case.get('witnesses', []),
                        'legal_persons': existing_case.get('legal_persons', [])
                    }
                    found = True
                    break

            if not found:
                # 新建案件
                case_data = {
                    'case_number': case_number,
                    'person_name': person_name,
                    'applicant': data.get('申请人', ''),
                    'case_type': data.get('案件类型', ''),
                    'year': datetime.now().year,
                    'folder_path': f"{datetime.now().year}/{case_number}",
                    'created_date': datetime.now().strftime('%Y-%m-%d'),
                    'employer': data.get('用人单位', ''),
                    'work_unit': data.get('用工单位', ''),
                    'workplace': data.get('工作场所', ''),
                    'regulation': data.get('条例', ''),
                    'operator': data.get('操作员', ''),
                    'person_info': {
                        'name': data.get('本人姓名', ''),
                        'gender': data.get('本人性别', ''),
                        'age': data.get('本人年龄', ''),
                        'phone': data.get('本人电话', ''),
                        'id_card': data.get('本人身份证号', ''),
                        'address': data.get('本人身份证地址', ''),
                        'current_address': data.get('本人现住址', ''),
                        'position': data.get('本人岗位', ''),
                        '自我介绍': data.get('自我介绍', ''),
                        '受伤经过': data.get('受伤经过', ''),
                        '就医情况': data.get('就医情况', ''),
                        '医疗结论': data.get('医疗结论', '')
                    },
                    'witnesses': [],
                    'legal_persons': []
                }
                index_data['cases'].append(case_data)

            # 更新统计信息
            index_data['total_cases'] = len(index_data['cases'])
            index_data['last_update'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            # 写入文件 - 只调用一次！
            success = self._write_case_index(index_data)
            if not success:
                print("警告：索引文件写入失败")

        except Exception as e:
            print(f"更新索引失败: {e}")
            traceback.print_exc()

    def get_current_year_folder(self):
        """获取当前年份的cases文件夹"""
        current_year = datetime.now().year
        year_folder = os.path.join(self.BASE_DIR, str(current_year))
        os.makedirs(year_folder, exist_ok=True)
        return year_folder

    def generate_case_number(self, injured_name):
        """生成案本号：类型-姓名-序号按年份"""
        # 确定类型前缀
        case_type = self.check_case_type()
        prefix = self.CASE_PREFIX_MAP.get(case_type, "GS")

        # 使用年份文件夹
        year_folder = self.get_current_year_folder()

        # 计算下一个序号
        existing_numbers = []
        if os.path.exists(year_folder):
            for folder in os.listdir(year_folder):
                # 匹配格式：前缀-姓名-数字
                if folder.startswith(f"{prefix}-{injured_name}-"):
                    try:
                        num = int(folder.split('-')[-1])
                        existing_numbers.append(num)
                    except:
                        continue

        # 生成新序号
        if existing_numbers:
            next_num = max(existing_numbers) + 1
        else:
            next_num = 1

        case_number = f"{prefix}-{injured_name}-{next_num:03d}"
        return case_number

    def closeEvent(self, event):
        """窗口关闭时最后保存一次"""
        if self.checkBox_remember.isChecked():
            operator = self.lineEdit_operator.text().strip()
            api_url = self.lineEdit_api_url.text().strip()
            api_key = self.lineEdit_api_key.text().strip()
            self.config.save_config(operator, api_url, api_key, True)

        # 保存窗口大小位置
        settings = QSettings("WorkInjuryApp", "Window")
        settings.setValue("geometry", self.saveGeometry())

        event.accept()

    # 以下是测试程序，编程完成以后需要删除
    def keyPressEvent(self, event):
        """键盘按下事件"""
        if event.key() == Qt.Key_F2:  # 按 F2 键
            self.fill_test_data()
        elif event.key() == Qt.Key_F3:  # 按 F3 键填下一组
            self.fill_next_test_data()

    def fill_test_data(self):
        """填入测试数据第一组"""
        self.test_index = getattr(self, 'test_index', 0)
        self.fill_next_test_data()

    def fill_next_test_data(self):
        """填入下一组测试数据"""
        # 测试数据10组
        test_data = [
            {
                "name": "张三",
                "id_card": "410101199001011234",
                "id_address": "河南省郑州市中原区建设路1号",
                "current_address": "河南省郑州市金水区花园路2号院3号楼",
                "phone": "13800138000",
                "position": "车间主任"
            },
            {
                "name": "李四",
                "id_card": "410101199105022345",
                "id_address": "河南省洛阳市西工区中州路5号",
                "current_address": "河南省洛阳市涧西区南昌路8号院",
                "phone": "13900139001",
                "position": "技术员"
            },
            {
                "name": "王五",
                "id_card": "410101198206033456",
                "id_address": "河南省开封市龙亭区中山路10号",
                "current_address": "河南省开封市禹王台区五一路3号",
                "phone": "13700137002",
                "position": "安全员"
            },
            {
                "name": "赵六",
                "id_card": "410101198503044567",
                "id_address": "河南省新乡市红旗区平原路15号",
                "current_address": "河南省新乡市卫滨区解放路20号院",
                "phone": "13600136003",
                "position": "操作工"
            },
            {
                "name": "孙七",
                "id_card": "410101197808055678",
                "id_address": "河南省焦作市山阳区塔南路25号",
                "current_address": "河南省焦作市中站区跃进路8号",
                "phone": "13500135004",
                "position": "维修工"
            },
            {
                "name": "周八",
                "id_card": "410101199212066789",
                "id_address": "河南省濮阳市华龙区中原路30号",
                "current_address": "河南省濮阳市濮阳县红旗路12号",
                "phone": "13400134005",
                "position": "电工"
            },
            {
                "name": "吴九",
                "id_card": "410101198909077890",
                "id_address": "河南省许昌市魏都区七一路18号",
                "current_address": "河南省许昌市建安区新许路6号",
                "phone": "13300133006",
                "position": "焊工"
            },
            {
                "name": "郑十",
                "id_card": "410101199311088901",
                "id_address": "河南省漯河市郾城区黄河路22号",
                "current_address": "河南省漯河市源汇区人民路15号",
                "phone": "13200132007",
                "position": "司机"
            },
            {
                "name": "钱多多",
                "id_card": "410101198012099012",
                "id_address": "河南省三门峡市湖滨区崤山路35号",
                "current_address": "河南省三门峡市陕州区神泉路9号",
                "phone": "13100131008",
                "position": "仓库管理员"
            },
            {
                "name": "刘能",
                "id_card": "410101199510101123",
                "id_address": "河南省南阳市卧龙区中州路45号",
                "current_address": "河南省南阳市宛城区建设路28号",
                "phone": "13000130009",
                "position": "质检员"
            }
        ]

        # 获取当前索引
        if not hasattr(self, 'test_index'):
            self.test_index = 0

        # 取当前组数据
        data = test_data[self.test_index]

        # 填入数据
        self.lineEdit_name.setText(data['name'])
        self.lineEdit_id_card.setText(data['id_card'])
        self.lineEdit_id_address.setText(data['id_address'])
        self.lineEdit_current_address.setText(data['current_address'])
        self.lineEdit_phone.setText(data['phone'])
        self.lineEdit_position.setText(data['position'])

        # 触发自动计算
        self.auto_calculate_id_info()

        # 更新索引循环
        self.test_index = (self.test_index + 1) % len(test_data)

        self.statusBar().showMessage(f"已填入测试数据 ({self.test_index}/{len(test_data)}): {data['name']}", 2000)

    def get_template_name(self, data):
        """根据人员类型和条例返回对应的模板文件名"""
        person_type = data['人员类型']  # 本人/证人/法人
        regulation = data.get('条例', '')  # 获取条例

        # 条例到文件名的映射
        regulation_map = {
            "第十四条第一款第一项（普通工伤案件）": "普通工伤案件",
            "第十四条第一款第二项（预备收尾案件）": "预备收尾案件",
            "第十四条第一款第三项（暴力伤害案件）": "暴力伤害案件",
            "第十四条第一款第四项（患职业病案件）": "患职业病案件",
            "第十四条第一款第五项（因工外出案件）": "因工外出案件",
            "第十四条第一款第六项（上下班时案件）": "上下班时案件",
            "第十五条第一款第一项（工作时因病亡故案件）": "工作时因病亡故案件",
            # 可以继续添加其他映射
        }

        # 获取条例对应的案件类型，如果没有匹配则用"普通工伤案件"
        case_type = regulation_map.get(regulation, "普通工伤案件")

        # 生成模板名：人员类型 + 谈话笔录（ + 案件类型 + ）.docx
        template_name = f"{person_type}谈话笔录（{case_type}）.docx"
        return template_name


def main():
    app = QApplication(sys.argv)

    app.setApplicationName("工伤案件管理系统")
    app.setOrganizationName("WorkInjuryApp")

    window = MainWindow()
    window.show()

    sys.exit(app.exec_())

if __name__ == "__main__":
    main()